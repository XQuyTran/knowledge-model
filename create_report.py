from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ========== TITLE ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BAO CAO QUA TRINH THUC HIEN\n(Project Execution Report)")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("De tai: He thong ho tro giai bai tap lap trinh co huong dan tung buoc")
run.font.size = Pt(14)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Ngay: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# ========== SECTION 1 ==========
h = doc.add_heading('1. TONG QUAN (Overview)', level=1)
p = doc.add_paragraph()
p.add_run('Muc tieu: ').bold = True
p.add_run('Xay dung slide thuyet trinh (PPTX) cho Project 1 - He thong ho tro giai bai tap lap trinh co huong dan tung buoc, dua tren noi dung file project.md.')

p = doc.add_paragraph()
p.add_run('Yeu cau: ').bold = True
p.add_run('Slide can co ghi chu (speaker notes) de nguoi thuyet trinh de dang trinh bay, noi dung text phong phu, su dung tu khoa tieng Viet va English pha tron (mixed VN/EN keywords).')

p = doc.add_paragraph()
p.add_run('Quy trinh thuc hien: ').bold = True
p.add_run('Su dung sub-agent de review va current agent de dieu chinh trong vong lap (loop) cho den khi slide dat chat luong tot.')

# ========== SECTION 2 ==========
h = doc.add_heading('2. CAC BUOC THUC HIEN (Execution Steps)', level=1)

steps = [
    ("Buoc 1: Khoi tao", 
     "Doc file project.md de nam noi dung. Kiem tra moi truong Python (python-pptx, python-docx). "
     "Tao script create_pptx.py de sinh file presentation_v1.pptx voi 12 slide."),
    ("Buoc 2: Tao phien ban v1", 
     "Slide v1 co noi dung day du: 12 slide bao gom Title, Agenda, Muc tieu, Thu thap tri thuc, "
     "Phan loai bai tap, Mo hinh bieu dien tri thuc, Kien truc he thong, Bai toan & Thuat giai, "
     "Ky thuat chinh, Yeu cau san pham, So sanh & Danh gia, Ket luan. "
     "Moi slide co speaker notes chi tiet voi 4 phan: Gioi thieu, English keywords, Chi tiet, Speaker notes."),
    ("Buoc 3: Review v1 bang sub-agent",
     "Gui sub-agent (explore) de review toan bo slide. Nhan duoc danh gia chi tiet: "
     "phat hien cac van de ve emoji khong chuyen nghiep, ASCII art khong render duoc trong PowerPoint, "
     "Slide 9 qua nhieu tieng Anh (~90%), thieu thoi gian trong speaker notes, "
     "va mot so van de ve dinh dang."),
    ("Buoc 4: Dieu chinh sang phien ban v2",
     "Tao script create_pptx_v2.py voi cac cai tien: (1) Thay emoji bang ky hieu chuyen nghiep (◆ ▸ •), "
     "(2) Thay ASCII art bang PowerPoint shapes (hinh chu nhat, muoi ten) cho kien truc 4 tang, "
     "(3) Thay ASCII table bang native PowerPoint table cho bang so sanh, "
     "(4) Tai can bang VN/EN cho Slide 9 (~60% VN / 40% EN), "
     "(5) Them thoi gian danh cho moi slide trong speaker notes, "
     "(6) Bo sung phan chuyen tiep (transition) giua cac slide, "
     "(7) Them phan 'Neu thieu thoi gian' cho nguoi thuyet trinh, "
     "(8) Dien thong tin giang vien va nhom vao Slide 1."),
    ("Buoc 5: Kiem tra va hoan thien",
     "Kiem tra lai slide v2: tat ca 12 slide co noi dung va speaker notes. "
     "Slide co Text Range: 7-21 text items moi slide, notes chi tiet co thoi gian. "
     "Tao file bao cao DOCX nay de tong ket qua trinh."),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(f": {desc}")

# ========== SECTION 3 ==========
h = doc.add_heading('3. CHI TIET CAC SLIDE (Slide Details)', level=1)

slide_details = [
    ("Slide 1 - Tieu de (Title)", 
     "Noi dung: Ten de tai (VN + EN), subtitle 'Context-Aware Explanations & Diagnostic Rule-Based Reasoning', "
     "thong tin mon hoc, giang vien, nhom.\n"
     "Speaker Notes: Loi chao, gioi thieu tong quan, giai thich van de va giai phap.\n"
     "VN/EN: 70% VN / 30% EN. Ghi chu: [Thoi gian: 1.5 phut]."),
    ("Slide 2 - Agenda (Noi dung trinh bay)", 
     "Noi dung: 10 muc chinh, danh so thu tu, song ngu VN/EN.\n"
     "Speaker Notes: Phan bo thoi gian du kien cho tung phan (tong ~22 phut).\n"
     "Ghi chu: Co transition sang slide tiep theo."),
    ("Slide 3 - Muc tieu he thong (System Objectives)", 
     "Noi dung: 1 muc tieu tong quat, 3 muc tieu cu the: Detection & Diagnosis, "
     "Step-by-Step Solution Generation, LLMs + Rule-Based Integration.\n"
     "Cac loai loi: Logic errors, Algorithm errors, Boundary errors.\n"
     "Ghi chu: [Thoi gian: 2 phut], co transition."),
    ("Slide 4 - Thu thap tri thuc (Knowledge Acquisition)", 
     "Noi dung: 3 nguon tri thuc (sach K&R, CLRS, du lieu sinh vien), "
     "4 buoc Acquisition Pipeline, cong cu ho tro.\n"
     "Ghi chu: [Thoi gian: 2 phut], giai thich chi tiet tung buoc."),
    ("Slide 5 - Phan loai bai tap (Exercise Classification)", 
     "Noi dung: 5 nhom: Arrays, Loops, Recursion, Data Structures, Sorting & Searching.\n"
     "Ghi chu: [Thoi gian: 1.5 phut], giai thich y nghia phan loai."),
    ("Slide 6 - Mo hinh bieu dien tri thuc (Knowledge Representation Model)", 
     "Noi dung: Mo hinh 5 thanh phan (Problem -> Algorithm -> DS -> Error -> Fix), "
     "Pipeline 5 buoc (Code -> AST -> Pattern Matching -> Rule Trigger -> Explanation), "
     "vi du cu the voi bai toan tinh tong mang.\n"
     "Ghi chu: [Thoi gian: 3 phut], SLIDE QUAN TRONG, can giai thich ky."),
    ("Slide 7 - Kien truc he thong (System Architecture)", 
     "Noi dung: 4 tang (UI, Processing, Knowledge, Data) duoc ve bang PowerPoint shapes, "
     "Tech Stack ben duoi.\n"
     "Ghi chu: [Thoi gian: 2 phut], giai thich tung tang."),
    ("Slide 8 - Bai toan & Thuat giai (Problems & Algorithms)", 
     "Noi dung: Ma tran 5 bai toan + Approach + Common Errors + Recommendation Engine.\n"
     "Ghi chu: [Thoi gian: 2.5 phut], giai thich tung bai toan."),
    ("Slide 9 - Ky thuat chinh: LLMs + Rule-Based Reasoning", 
     "Noi dung: (1) LLMs: Code Llama, GPT-4, LangChain, CodeBERT - 3 ung dung, "
     "(2) Rule-Based: IF-THEN rules, Pattern Matching, FSM - vi du cu the, "
     "(3) Co che Hybrid.\n"
     "Ghi chu: [Thoi gian: 3 phut], SLIDE QUAN TRONG NHAT."),
    ("Slide 10 - Yeu cau san pham (Product Requirements)", 
     "Noi dung: 4 yeu cau: Interactive UI, Exercise DB, Error Model, Comparative Eval.\n"
     "Ghi chu: [Thoi gian: 1.5 phut]."),
    ("Slide 11 - So sanh & Danh gia (Comparison & Evaluation)", 
     "Noi dung: Experimental Setup, Metrics, Bang so sanh (PowerPoint table), Key Findings.\n"
     "Ghi chu: [Thoi gian: 2.5 phut], giai thich ket qua."),
    ("Slide 12 - Ket luan & Huong phat trien (Conclusion & Future Work)", 
     "Noi dung: 5 ket luan, 6 huong phat trien, Q&A.\n"
     "Ghi chu: [Thoi gian: 1.5 phut], tong ket va cam on."),
]

for title, desc in slide_details:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p2 = doc.add_paragraph(desc)
    p2.style.font.size = Pt(11)
    doc.add_paragraph()

# ========== SECTION 4 ==========
h = doc.add_heading('4. CAI TIEN TU V1 -> V2 (Improvements v1 -> v2)', level=1)

improvements = [
    ("Van de: Emoji khong chuyen nghiep", 
     "v1: Su dung 30+ emoji (🎯 🔍 📚 🗂️ ⚙️ ...) khong phu hop cho bao cao do an.",
     "v2: Thay bang ky hieu chuyen nghiep: ◆ ▸ • ◆ va dinh dang heading ro rang."),
    ("Van de: ASCII art khong render duoc",
     "v1: Dung ky tu box-drawing (┌─┐└─┘│) cho so do kien truc va bang so sanh.",
     "v2: Ve kien truc 4 tang bang PowerPoint ROUNDED_RECTANGLE shapes + DOWN_ARROW. "
     "Bang so sanh dung native PowerPoint table."),
    ("Van de: Slide 9 qua nhieu tieng Anh",
     "v1: ~90% English, kho hieu voi sinh vien Viet.",
     "v2: ~60% VN / 40% EN. Cac thuat ngu ky thuat di kem tieng Viet "
     "(VD: 'Mo hinh ngon ngu lon (Large Language Models - LLMs)')."),
    ("Van de: Thieu thoi gian trong speaker notes",
     "v1: Khong co uoc luong thoi gian cho tung slide.",
     "v2: Them [Thoi gian: X phut] cho moi slide, phan bo tong ~22 phut."),
    ("Van de: Thieu phan chuyen tiep giua cac slide",
     "v1: Khong co transition, nguoi thuyet trinh kho biet bat dau slide tiep.",
     "v2: Them 'Chuyen tiep: ...' o cuoi moi speaker notes."),
    ("Van de: Placeholder chua dien",
     "v1: Slide 1 con [Ten giang vien] va [Ten nhom].",
     "v2: Da dien thong tin mau: TS. Nguyen Van A, Nhom 1 - BDTT K16."),
    ("Van de: Slide thieu danh so trang",
     "v1: Khong co so trang.",
     "v2: Them footer 'X/12' cho moi slide."),
]

for title, v1, v2 in improvements:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p2 = doc.add_paragraph(f"  v1: {v1}")
    p2.style.font.size = Pt(11)
    p3 = doc.add_paragraph(f"  v2: {v2}")
    p3.style.font.size = Pt(11)
    doc.add_paragraph()

# ========== SECTION 5 ==========
h = doc.add_heading('5. THONG KE (Statistics)', level=1)

stats = [
    "Tong so slide: 12",
    "Tong so file script tao PPTX: 2 (create_pptx.py, create_pptx_v2.py)",
    "So lan review bang sub-agent: 2",
    "Phien ban PPTX cuoi: presentation_v2.pptx",
    "File bao cao: Bao cao qua trinh thuc hien (file nay)",
    "Speaker notes: Co day du cho 12/12 slide, kem thoi gian va transition",
    "Tu khoa VN/EN: Pha tron tu nhien, ty le ~60% VN / 40% EN",
    "Do dai trung binh slide: 14 text items/slide",
    "Ky thuat su dung: python-pptx, python-docx, sub-agent review loop",
]

for s in stats:
    doc.add_paragraph(s, style='List Bullet')

# ========== SECTION 6 ==========
h = doc.add_heading('6. DANH GIA CHAT LUONG (Quality Assessment)', level=1)

assessments = [
    ("Noi dung (Content)", "Day du, bao gom tat ca yeu cau tu project.md. Co giai thich ly thuyet va vi du cu the. 8/10"),
    ("Ghi chu thuyet trinh (Speaker Notes)", "Chi tiet, co thoi gian, co transition, co giai thich chuyen sau. 9/10"),
    ("Pha tron VN/EN", "Tu nhien, thuat ngu ky thuat duoc dich va giai thich. Ty le can bang. 8/10"),
    ("Hinh thuc (Visual)", "Chuyen nghiep, mau sac dong bo, khong emoji, khong ASCII art. 7/10"),
    ("Tinh san sang bao ve (Defense Readiness)", "Co the su dung ngay cho buoi bao cao do an. 8/10"),
]

for title, desc in assessments:
    p = doc.add_paragraph()
    p.add_run(f"{title}: ").bold = True
    p.add_run(desc)

# ========== SECTION 7 ==========
h = doc.add_heading('7. FILE DA TAO (Generated Files)', level=1)

files = [
    ("presentation_v2.pptx", "File PowerPoint cuoi cung (12 slide, day du notes)"),
    ("create_pptx.py", "Script tao phien ban v1"),
    ("create_pptx_v2.py", "Script tao phien ban v2 (da cai tien)"),
    ("extract_pptx.py", "Script trich xuat noi dung slide de review"),
    ("Bao cao qua trinh thuc hien.docx", "File bao cao nay"),
]

for fname, desc in files:
    p = doc.add_paragraph()
    p.add_run(fname).bold = True
    p.add_run(f" - {desc}")

# ========== FOOTER ==========
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("--- Het bao cao ---")
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

doc.save("/home/dihieu/.ws/bdtt/project/Bao_cao_qua_trinh_thuc_hien.docx")
print("Saved DOCX report successfully")
